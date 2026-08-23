import re
import io
from pypdf import PdfReader
from docx import Document

def extract_text_from_pdf(file_bytes_or_path):
    """Extract text from PDF file stream or file path."""
    text = ""
    try:
        if isinstance(file_bytes_or_path, bytes):
            reader = PdfReader(io.BytesIO(file_bytes_or_path))
        else:
            reader = PdfReader(file_bytes_or_path)
        
        for page in reader.pages:
            extracted = page.extract_text()
            if extracted:
                text += extracted + "\n"
    except Exception as e:
        print(f"Error reading PDF: {e}")
        try:
            from pdfminer.high_level import extract_text as pdfminer_extract
            if isinstance(file_bytes_or_path, bytes):
                text = pdfminer_extract(io.BytesIO(file_bytes_or_path))
            else:
                text = pdfminer_extract(file_bytes_or_path)
        except Exception as e2:
            print(f"Pdfminer fallback error: {e2}")
    return text.strip()

def extract_text_from_docx(file_bytes_or_path):
    """Extract text from DOCX file stream or file path."""
    text = ""
    try:
        if isinstance(file_bytes_or_path, bytes):
            doc = Document(io.BytesIO(file_bytes_or_path))
        else:
            doc = Document(file_bytes_or_path)
            
        full_text = []
        for para in doc.paragraphs:
            if para.text:
                full_text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        full_text.append(cell.text)
        text = "\n".join(full_text)
    except Exception as e:
        print(f"Error reading DOCX: {e}")
    return text.strip()

def validate_resume_structure(text):
    """
    Checks whether uploaded document contains standard resume section indicators.
    Returns tuple (is_resume: bool, warning_message: str).
    """
    if not text or len(text.strip()) < 40:
        return False, "⚠️ Non-Resume Document / Unreadable File Content"
        
    text_lower = text.lower()
    resume_keywords = ['experience', 'education', 'skills', 'projects', 'summary', 'employment', 'curriculum', 'cv', 'work history', 'contact']
    matches = sum(1 for kw in resume_keywords if kw in text_lower)
    
    if matches < 2:
        return False, "⚠️ Potential Non-Resume File (Lacks Standard Resume Sections)"
    return True, "Valid Resume Format"

def parse_resume_file(uploaded_file):
    """
    Accepts Streamlit UploadedFile object or filepath tuple.
    Returns dict with extracted text, contact details, experience, and document validity status.
    """
    if hasattr(uploaded_file, "name"):
        filename = uploaded_file.name
        content = uploaded_file.getvalue()
    elif isinstance(uploaded_file, tuple):
        filename, content = uploaded_file
    elif isinstance(uploaded_file, str):
        filename = uploaded_file
        with open(uploaded_file, "rb") as f:
            content = f.read()
    else:
        filename = "document.txt"
        content = str(uploaded_file).encode('utf-8')

    ext = filename.lower().split('.')[-1]
    raw_text = ""

    if ext == 'pdf':
        raw_text = extract_text_from_pdf(content)
    elif ext in ['docx', 'doc']:
        raw_text = extract_text_from_docx(content)
    else:
        if isinstance(content, bytes):
            try:
                raw_text = content.decode('utf-8')
            except Exception:
                raw_text = content.decode('latin-1', errors='ignore')
        else:
            raw_text = str(content)

    email = extract_email(raw_text)
    phone = extract_phone(raw_text)
    candidate_name = extract_candidate_name(raw_text, filename)
    exp_years = extract_experience_years(raw_text)
    education = extract_education_level(raw_text)
    is_valid_resume, doc_warning = validate_resume_structure(raw_text)

    return {
        "filename": filename,
        "text": raw_text,
        "candidate_name": candidate_name,
        "email": email,
        "phone": phone,
        "experience_years": exp_years,
        "education": education,
        "is_valid_resume": is_valid_resume,
        "doc_warning": doc_warning
    }

def extract_email(text):
    email_pattern = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
    matches = re.findall(email_pattern, text)
    return matches[0] if matches else "N/A"

def extract_phone(text):
    phone_pattern = r'(?:\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}'
    matches = re.findall(phone_pattern, text)
    return matches[0] if matches else "N/A"

def extract_candidate_name(text, filename):
    lines = [line.strip() for line in text.split('\n') if line.strip()]
    if lines:
        first_line = lines[0]
        if len(first_line) < 40 and not any(kw in first_line.lower() for kw in ['resume', 'curriculum', 'page', 'email', 'http', '@', 'chapter', 'lecture']):
            return first_line.title()
    
    clean_name = re.sub(r'[_#-]', ' ', filename.split('.')[0])
    clean_name = re.sub(r'(?i)(resume|cv|jd|job|description|senior|junior|lead)', '', clean_name).strip()
    return clean_name.title() if clean_name else "Candidate"

def extract_experience_years(text):
    patterns = [
        r'(\d+)\+?\s*(?:years?|yrs)\s*(?:of)?\s*(?:experience|exp)?',
        r'(?:experience|exp)\s*:\s*(\d+)\+?\s*(?:years?|yrs)',
        r'(\d+)\+?\s*years?\s*in\b'
    ]
    for pattern in patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        if matches:
            try:
                years = [int(m) for m in matches if int(m) <= 40]
                if years:
                    return max(years)
            except ValueError:
                pass

    year_matches = re.findall(r'\b(20\d{2}|19\d{2})\b', text)
    if year_matches:
        years = sorted([int(y) for y in year_matches if 1990 <= int(y) <= 2026])
        if len(years) >= 2:
            span = max(years) - min(years)
            if 0 < span <= 30:
                return span
    return 1

def extract_education_level(text):
    text_lower = text.lower()
    degrees = []
    if any(k in text_lower for k in ['ph.d', 'phd', 'doctorate']):
        degrees.append("Ph.D. / Doctorate")
    if any(k in text_lower for k in ['master', 'm.s.', 'm.tech', 'm.sc', 'mca', 'mba']):
        degrees.append("Master's Degree (M.S. / M.Tech)")
    if any(k in text_lower for k in ['bachelor', 'b.s.', 'b.tech', 'b.e.', 'b.sc', 'bca', 'bba']):
        degrees.append("Bachelor's Degree (B.S. / B.Tech)")
    if any(k in text_lower for k in ['associate', 'diploma']):
        degrees.append("Associate / Diploma")

    if degrees:
        return degrees[0]
    return "Bachelor's Degree (Equivalent)"
