import os
from fpdf import FPDF
from docx import Document

def create_sample_files():
    os.makedirs("sample_data/resumes", exist_ok=True)
    os.makedirs("sample_data/job_descriptions", exist_ok=True)

    # 1. Job Descriptions
    jd_ai_engineer = """
Job Title: Senior AI / Machine Learning Engineer
Department: Artificial Intelligence & Data Science
Location: Remote / San Francisco, CA
Experience Required: 3+ Years
Education: Bachelor's or Master's in Computer Science, AI, Data Science, or related field

Job Overview:
We are seeking an experienced AI/ML Engineer to build state-of-the-art NLP and machine learning models. You will be responsible for end-to-end model development, deployment, and optimization.

Key Responsibilities:
- Design, train, and deploy machine learning models using Python, PyTorch, or TensorFlow.
- Process large unstructured text datasets using Natural Language Processing (NLP) techniques.
- Build RESTful APIs using FastAPI or Flask to serve AI models in production.
- Containerize applications using Docker and deploy on AWS or Cloud infrastructure.
- Implement CI/CD pipelines and monitor model performance in production.
- Collaborate with cross-functional teams to integrate AI capabilities into product features.

Required Skills & Qualifications:
- Technical Skills: Python, Machine Learning, Deep Learning, PyTorch, TensorFlow, NLP, Scikit-Learn, Pandas, NumPy, SQL, Docker, FastAPI, AWS, Git.
- Experience: Minimum 3 years of hands-on experience in AI/ML software development.
- Education: B.Tech / B.S. or M.Tech / M.S. in Computer Science or Data Science.
- Soft Skills: Problem solving, Communication, Teamwork, Analytical thinking.
"""

    jd_fullstack = """
Job Title: Full Stack Web Developer
Department: Engineering
Location: Hybrid / New York, NY
Experience Required: 2+ Years
Education: Bachelor's degree in Computer Science, Software Engineering, or equivalent experience

Job Overview:
We are looking for a talented Full Stack Web Developer to design and implement robust, scalable web applications. You will work across the entire stack, from modern frontend interfaces to backend microservices.

Key Responsibilities:
- Develop responsive, user-friendly frontend web interfaces using React, JavaScript, HTML5, and CSS3.
- Build performant RESTful APIs and backend services using Python, Node.js, and Express.
- Manage relational and NoSQL databases including PostgreSQL, MySQL, and MongoDB.
- Write unit tests and maintain continuous integration and deployment pipelines.
- Optimize web applications for maximum speed and scalability.

Required Skills & Qualifications:
- Technical Skills: Python, React, JavaScript, Node.js, Express, HTML5, CSS3, PostgreSQL, MongoDB, REST API, Git, Docker, TypeScript.
- Experience: 2+ years of full stack web development experience.
- Soft Skills: Agile methodology, Collaboration, Code review, Problem solving.
"""

    jd_data_analyst = """
Job Title: Data Analyst & Business Intelligence Specialist
Department: Data Analytics
Location: Remote
Experience Required: 1+ Years
Education: Bachelor's degree in Statistics, Computer Science, Economics, or Data Analytics

Job Overview:
We are hiring a Data Analyst to transform complex data into actionable insights for business decision-making. You will create interactive dashboards, write SQL queries, and perform statistical data analysis.

Key Responsibilities:
- Query and analyze large datasets using SQL and Python (Pandas, NumPy).
- Create interactive dashboards and visual reports using Tableau or PowerBI.
- Conduct statistical analysis and data cleaning to identify trends and metrics.
- Present findings to leadership and stakeholders in clear, actionable visual reports.

Required Skills & Qualifications:
- Technical Skills: SQL, Python, Pandas, NumPy, Tableau, Power BI, Excel, Data Visualization, Statistics, Data Cleaning, ETL.
- Experience: 1+ years in data analysis, reporting, or business intelligence.
- Soft Skills: Data storytelling, Presentation skills, Attention to detail.
"""

    with open("sample_data/job_descriptions/AI_ML_Engineer_JD.txt", "w", encoding="utf-8") as f:
        f.write(jd_ai_engineer)
    with open("sample_data/job_descriptions/Full_Stack_Developer_JD.txt", "w", encoding="utf-8") as f:
        f.write(jd_fullstack)
    with open("sample_data/job_descriptions/Data_Analyst_JD.txt", "w", encoding="utf-8") as f:
        f.write(jd_data_analyst)

    # 2. Resumes
    resumes_data = [
        {
            "filename": "Alex_Rivera_Senior_AI_Engineer.pdf",
            "name": "Alex Rivera",
            "title": "Senior AI / Machine Learning Engineer",
            "email": "alex.rivera@email.com",
            "phone": "+1 (555) 234-5678",
            "location": "San Francisco, CA",
            "experience_years": "5 Years Experience",
            "education": "Master of Science (M.S.) in Computer Science - Stanford University (2019)",
            "summary": "Innovative Senior AI Engineer with 5+ years of experience designing and deploying production NLP and deep learning models. Expert in PyTorch, Python, Transformers, Docker, and AWS cloud deployment.",
            "skills": [
                "Python", "PyTorch", "TensorFlow", "NLP", "Machine Learning", "Deep Learning",
                "Scikit-Learn", "Pandas", "NumPy", "SQL", "FastAPI", "Docker", "AWS", "Git",
                "BERT", "Transformers", "CI/CD", "MLOps", "REST APIs"
            ],
            "experience": [
                "Lead AI Engineer at TechCore Inc. (2021 - Present): Developed LLM and NLP screening pipelines processing 100k+ documents daily. Deployed microservices using FastAPI, Docker, and AWS SageMaker.",
                "Machine Learning Developer at DataVision (2019 - 2021): Trained PyTorch deep learning models for text classification and NER extraction, improving accuracy by 18%."
            ]
        },
        {
            "filename": "Priya_Sharma_ML_Developer.pdf",
            "name": "Priya Sharma",
            "title": "Machine Learning Engineer",
            "email": "priya.sharma@email.com",
            "phone": "+1 (555) 876-5432",
            "location": "Seattle, WA",
            "experience_years": "3 Years Experience",
            "education": "Bachelor of Technology (B.Tech) in Computer Engineering - IIT Delhi (2021)",
            "summary": "Dedicated Machine Learning Engineer with 3 years of hands-on experience building NLP text analytics engines, predictive models, and REST microservices in Python.",
            "skills": [
                "Python", "Machine Learning", "Scikit-Learn", "NLP", "Pandas", "NumPy",
                "SQL", "Flask", "Docker", "Git", "PostgreSQL", "Data Science", "Natural Language Processing"
            ],
            "experience": [
                "ML Engineer at Nexus Systems (2022 - Present): Designed NLP keyword extraction and text classification workflows using Scikit-Learn and SpaCy.",
                "Data Science Associate at Analytics Hub (2021 - 2022): Built SQL pipelines and automated feature extraction scripts in Python."
            ]
        },
        {
            "filename": "Marcus_Chen_Data_Scientist.pdf",
            "name": "Marcus Chen",
            "title": "Data Scientist & Analytics Specialist",
            "email": "marcus.chen@email.com",
            "phone": "+1 (555) 345-6789",
            "location": "Austin, TX",
            "experience_years": "2 Years Experience",
            "education": "Bachelor of Science (B.S.) in Data Science - UT Austin (2022)",
            "summary": "Data Scientist passionate about extracting business insights from complex datasets using Python, SQL, statistical modeling, and Tableau visualization.",
            "skills": [
                "Python", "SQL", "Pandas", "NumPy", "Tableau", "Power BI", "Statistics",
                "Data Visualization", "Data Cleaning", "R", "Excel", "Scikit-Learn"
            ],
            "experience": [
                "Data Scientist at Horizon Analytics (2022 - Present): Created automated BI dashboards in Tableau and conducted statistical modeling in Python for client growth strategies."
            ]
        },
        {
            "filename": "Sophia_Taylor_FullStack_Dev.docx",
            "name": "Sophia Taylor",
            "title": "Senior Full Stack Web Developer",
            "email": "sophia.taylor@email.com",
            "phone": "+1 (555) 456-7890",
            "location": "New York, NY",
            "experience_years": "4 Years Experience",
            "education": "B.S. in Software Engineering - Columbia University (2020)",
            "summary": "Creative Full Stack Developer with 4 years of experience building modern web applications using React, Node.js, Python, TypeScript, and Docker.",
            "skills": [
                "Python", "React", "JavaScript", "TypeScript", "Node.js", "Express",
                "HTML5", "CSS3", "PostgreSQL", "MongoDB", "REST API", "Docker", "Git"
            ],
            "experience": [
                "Full Stack Developer at WebSphere Tech (2021 - Present): Developed high-traffic web applications with React frontend and FastAPI backend.",
                "Frontend Engineer at CloudApps (2020 - 2021): Built interactive React components and RESTful API integrations."
            ]
        },
        {
            "filename": "David_Miller_Junior_Developer.docx",
            "name": "David Miller",
            "title": "Junior Web Developer",
            "email": "david.miller@email.com",
            "phone": "+1 (555) 567-8901",
            "location": "Chicago, IL",
            "experience_years": "1 Year Experience",
            "education": "Bachelor of Science in Information Technology (2023)",
            "summary": "Enthusiastic Junior Web Developer with 1 year of experience creating responsive landing pages and basic web utilities.",
            "skills": [
                "HTML5", "CSS3", "JavaScript", "Python", "Git", "Bootstrap", "SQL"
            ],
            "experience": [
                "Junior Web Intern at DigitalCraft (2023 - Present): Assisted in front-end HTML/CSS layouts and simple Python scripting."
            ]
        },
        {
            "filename": "Emily_Watson_Marketing_Manager.pdf",
            "name": "Emily Watson",
            "title": "Digital Marketing & Brand Specialist",
            "email": "emily.watson@email.com",
            "phone": "+1 (555) 678-9012",
            "location": "Boston, MA",
            "experience_years": "4 Years Experience",
            "education": "B.A. in Communications & Marketing - Boston University (2020)",
            "summary": "Results-driven Digital Marketing Lead specializing in social media campaigns, SEO content creation, brand growth, and public relations.",
            "skills": [
                "Digital Marketing", "SEO", "Content Strategy", "Social Media", "Copywriting",
                "Google Analytics", "Brand Management", "Public Relations"
            ],
            "experience": [
                "Marketing Manager at MediaPro (2021 - Present): Led digital ad strategy and social media content campaigns across 5 major brands."
            ]
        }
    ]

    for cand in resumes_data:
        file_path = os.path.join("sample_data/resumes", cand["filename"])
        if cand["filename"].endswith(".pdf"):
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Helvetica", "B", 18)
            pdf.cell(0, 10, cand["name"], new_x="LMARGIN", new_y="NEXT")
            pdf.set_font("Helvetica", "I", 12)
            pdf.cell(0, 8, f"{cand['title']} | {cand['experience_years']}", new_x="LMARGIN", new_y="NEXT")
            pdf.set_font("Helvetica", "", 10)
            pdf.cell(0, 6, f"Email: {cand['email']} | Phone: {cand['phone']} | Location: {cand['location']}", new_x="LMARGIN", new_y="NEXT")
            pdf.ln(4)
            
            pdf.set_font("Helvetica", "B", 12)
            pdf.cell(0, 8, "SUMMARY", new_x="LMARGIN", new_y="NEXT")
            pdf.set_font("Helvetica", "", 10)
            pdf.multi_cell(0, 5, cand["summary"])
            pdf.ln(3)

            pdf.set_font("Helvetica", "B", 12)
            pdf.cell(0, 8, "EDUCATION", new_x="LMARGIN", new_y="NEXT")
            pdf.set_font("Helvetica", "", 10)
            pdf.cell(0, 6, cand["education"], new_x="LMARGIN", new_y="NEXT")
            pdf.ln(3)

            pdf.set_font("Helvetica", "B", 12)
            pdf.cell(0, 8, "SKILLS", new_x="LMARGIN", new_y="NEXT")
            pdf.set_font("Helvetica", "", 10)
            pdf.multi_cell(0, 5, ", ".join(cand["skills"]))
            pdf.ln(3)

            pdf.set_font("Helvetica", "B", 12)
            pdf.cell(0, 8, "WORK EXPERIENCE", new_x="LMARGIN", new_y="NEXT")
            pdf.set_font("Helvetica", "", 10)
            for exp in cand["experience"]:
                pdf.multi_cell(0, 5, f"- {exp}")
                pdf.ln(2)

            pdf.output(file_path)

        elif cand["filename"].endswith(".docx"):
            doc = Document()
            doc.add_heading(cand["name"], level=0)
            p = doc.add_paragraph()
            p.add_run(f"{cand['title']} | {cand['experience_years']}\n").italic = True
            p.add_run(f"Email: {cand['email']} | Phone: {cand['phone']} | Location: {cand['location']}")

            doc.add_heading("SUMMARY", level=1)
            doc.add_paragraph(cand["summary"])

            doc.add_heading("EDUCATION", level=1)
            doc.add_paragraph(cand["education"])

            doc.add_heading("SKILLS", level=1)
            doc.add_paragraph(", ".join(cand["skills"]))

            doc.add_heading("WORK EXPERIENCE", level=1)
            for exp in cand["experience"]:
                doc.add_paragraph(exp, style='List Bullet')

            doc.save(file_path)

    print("Sample JDs and Resumes generated successfully!")

if __name__ == "__main__":
    create_sample_files()
